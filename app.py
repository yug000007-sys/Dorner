import io
import os
import re
import zipfile
import html
from pathlib import Path
from datetime import datetime, timezone
from email.utils import parsedate_to_datetime
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st
from bs4 import BeautifulSoup
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import extract_msg
except Exception:
    extract_msg = None

APP_TITLE = "Dorner Lead Automation"
EASTERN = ZoneInfo("America/New_York")

HEADERS = [
    "Referral", "ReferralEmail", "Brand", "ReceivedDateTime", "FirstName", "LastName", "ContactTitle", "Email",
    "Company", "Address", "County", "City", "State", "ZipCode", "Country", "LeadSource1", "LeadSource2",
    "LeadSource3", "LeadSource4", "LeadComments", "PhoneSupplied", "PhSuppliedExtension", "PhoneResearched",
    "CSRName", "PDF", "DUNS", "WebAddress", "SIC", "NAICS", "noOfEmployees", "ParentName", "LineOfBusiness",
    "Product", "Market", "PQ", "interestedIn", "crm_lead_id", "Latitude", "Longitude", "Keyword", "device",
    "DemoLead", "about_me", "college_1", "college_1_degree", "college_1_start", "college_1_end", "college_2",
    "college_2_degree", "college_2_start", "college_2_end", "month_of_joining", "about_experience", "Linkedin_Link",
    "Linkedin_Title", "searched_on_google", "linkedin_city", "linkedin_state", "linkedin_country", "GrandTotal"
]

CAD_COMMENT = "Please find the following new RFQ and URGENT lead from Dorner CAD and process accordingly. Kindly contact the customer to review the application to ensure the proper equipment is quoted based upon the application requirements. Please click on 'Click Here' below to view the lead details."
CONFIG_COMMENT = "Please find the following new RFQ and URGENT lead from Dorner Config and process accordingly. Kindly contact the customer to review the application to ensure the proper equipment is quoted based upon the application requirements. Please click on 'Click Here' below to view the lead details."


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"<mailto:([^>]+)>", "", text)
    text = re.sub(r"<https?://[^>]+>", "", text)
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def html_to_text(html_body: str) -> str:
    soup = BeautifulSoup(html_body or "", "lxml")
    for a in soup.find_all("a"):
        a.replace_with(a.get_text(" "))
    text = soup.get_text("\n")
    lines = [ln.strip() for ln in text.splitlines()]
    return clean_text("\n".join([ln for ln in lines if ln]))


def parse_msg_bytes(data: bytes, filename: str):
    tmp = Path("/tmp") / filename
    tmp.write_bytes(data)
    subject = Path(filename).stem
    body = ""
    raw_date = None
    sender = ""

    if extract_msg and filename.lower().endswith(".msg"):
        try:
            msg = extract_msg.Message(str(tmp))
            subject = msg.subject or subject
            sender = msg.sender or ""
            # Prefer the actual delivery/received timestamp over the "Date:" header
            # (which reflects when the sender's client composed/sent the message).
            # Both are real datetime objects from extract_msg, not strings -- keep
            # them as datetimes so normalize_datetime doesn't have to re-parse a
            # stringified ISO value with an RFC2822 parser (which silently fails).
            raw_date = getattr(msg, "receivedTime", None) or getattr(msg, "date", None)
            body = msg.body or ""
            if not body and getattr(msg, "htmlBody", None):
                body = html_to_text(msg.htmlBody)
            msg.close()
        except Exception:
            body = data.decode("utf-8", errors="ignore")
    else:
        body = data.decode("utf-8", errors="ignore")
        m = re.search(r"^Subject:\s*(.+)$", body, re.I | re.M)
        if m:
            subject = m.group(1).strip()
        m = re.search(r"^Date:\s*(.+)$", body, re.I | re.M)
        if m:
            raw_date = m.group(1).strip()

    body = clean_text(body)
    return subject, body, raw_date, sender


def first_match(text, patterns, default=""):
    for pat in patterns:
        m = re.search(pat, text, re.I | re.S)
        if m:
            return m.group(1).strip()
    return default


def normalize_datetime(raw_date, body: str) -> tuple[str, str]:
    candidates = []
    if raw_date:
        candidates.append(raw_date)
    created = first_match(body, [r"Created:\s*(.+)"])
    if created:
        candidates.append(created)

    for c in candidates:
        dt = None

        # extract_msg's .date / .receivedTime are already real datetime objects --
        # use them directly instead of stringifying and re-parsing (which fails
        # silently since str(datetime) is ISO format, not RFC2822).
        if isinstance(c, datetime):
            dt = c
        else:
            c = str(c).strip()
            try:
                dt = parsedate_to_datetime(c)
            except Exception:
                dt = None
            if dt is None:
                for fmt in ["%a %d %b %Y %I:%M:%S %p %Z", "%a %d %b %Y %H:%M:%S %Z", "%m/%d/%Y %I:%M %p"]:
                    try:
                        dt = datetime.strptime(c.replace(" UTC", " GMT"), fmt.replace("UTC", "GMT"))
                        break
                    except Exception:
                        continue

        if dt is None:
            continue

        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        et = dt.astimezone(EASTERN)
        return et.strftime("%-m/%-d/%Y %-I:%M %p"), et.strftime("%Y%m%d_%H%M%S")

    now = datetime.now(EASTERN)
    return now.strftime("%-m/%-d/%Y %-I:%M %p"), now.strftime("%Y%m%d_%H%M%S")


def parse_address(block: str):
    address = first_match(block, [r"Address:\s*(.*?)(?:\n\s*\n|Dorner Quote:|CAD Models|$)"])
    lines = [x.strip() for x in address.splitlines() if x.strip()]
    street = lines[0] if lines else ""
    city = state = zip_code = country = ""
    if len(lines) >= 2:
        m = re.search(r"(.+?)\s+([A-Z]{2})\s+([A-Z0-9\- ]{4,10})\s+([A-Z]{2,3})$", lines[-1])
        if m:
            city, state, zip_code, country = [g.strip() for g in m.groups()]
    return street, city, state, zip_code, country


def extract_device(body: str) -> str:
    start = body.find("Distributor:")
    if start < 0:
        start = body.find("Customer Contact Info:")
    if start < 0:
        start = 0
    device = body[start:].strip()
    device = re.sub(r"<https?://[^>]+>", "", device)
    device = re.sub(r"https?://\S+", "", device)
    device = re.sub(r"\n{4,}", "\n\n\n", device)
    return device.strip()


def extract_grand_total(text: str) -> str:
    matches = re.findall(r"Grand\s+Total:\s*\$?\s*([0-9,]+(?:\.\d{2})?)", text, flags=re.I)
    if matches:
        return "$" + matches[-1]
    return ""


def determine_brand_product(text: str):
    low = text.lower()
    if "garvey" in low:
        return "Garvey", ""
    if "montratec" in low:
        return "Montratec", ""
    if "aquagard" in low or "aquapruf" in low:
        return "Dorner", "AquaX"
    return "Dorner", ""


def parse_lead(subject: str, body: str, raw_date: str, original_filename: str):
    dt_text, stamp = normalize_datetime(raw_date, body)
    file_base = f"Dorner_{stamp}"
    device = extract_device(body)
    # [ \t]* (not \s*) confines the match to the SAME line as the label. \s*
    # includes \n, so when a field is genuinely blank in the source, a greedy
    # \s* silently skips every blank line and captures the next section's
    # header text instead (e.g. Distributor becomes "Customer Contact Info:").
    # [^\n]* (not [^\n]+) lets a truly blank field match as empty right there,
    # instead of failing to match and leaving the field to whatever the next
    # occurrence of similar text happens to be.
    name = first_match(body, [r"Name:[ \t]*([^\n]*)"])
    parts = name.split()
    first = parts[0] if parts else ""
    last = " ".join(parts[1:]) if len(parts) > 1 else ""
    street, city, state, zip_code, country = parse_address(body)
    brand, product = determine_brand_product(device)
    lead_comment = CAD_COMMENT if "Dorner CAD" in body else CONFIG_COMMENT if "Dorner Config" in body else ""
    pdf_names = f"{file_base}.pdf, {file_base}.msg, {file_base}.docx"

    row = {h: "" for h in HEADERS}
    row.update({
        "Brand": brand,
        "Product": product,
        "ReceivedDateTime": dt_text,
        "FirstName": first,
        "LastName": last,
        "ContactTitle": first_match(body, [r"Title:[ \t]*([^\n]*)"]),
        "Email": first_match(body, [r"Email:[ \t]*([^\s\n]*)"]),
        "Company": first_match(body, [r"Company:[ \t]*([^\n]*)"]),
        "Address": street,
        "City": city,
        "State": state,
        "ZipCode": zip_code,
        "Country": country,
        "LeadSource1": "Request For Quote",
        "LeadComments": lead_comment,
        "PhoneSupplied": first_match(body, [r"Phone:[ \t]*([^\n]*)"]),
        "PDF": pdf_names,
        "Keyword": subject,
        "device": device,
        "GrandTotal": extract_grand_total(body),
    })
    row["_file_base"] = file_base
    row["_original_filename"] = original_filename
    row["_quote"] = first_match(body, [r"Dorner Quote:\s*([0-9]+)", r"Quote\s+([0-9]{5,})"])
    row["_lead_time"] = first_match(body, [r"Lead Time \(Business Days\)\s*([0-9]+)"])
    return row


ORANGE_FILL = "FF6600"
LIGHT_BLUE_FILL = "C6D9F1"

# OOXML requires tblPr/tcPr children in a strict schema-defined sequence.
# Blindly .append()-ing a new element (e.g. tblBorders after python-docx has
# already added tblW/tblLayout/tblLook) produces a file that LibreOffice opens
# without complaint but that fails XSD validation and Word can reject as
# corrupt. Insert at the schema-correct position instead of appending.
_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
    "tblDescription",
]
_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
]


def _insert_ordered(parent, new_el, order):
    new_tag = qn_local(new_el.tag)
    new_idx = order.index(new_tag) if new_tag in order else len(order)
    for child in parent:
        child_tag = qn_local(child.tag)
        child_idx = order.index(child_tag) if child_tag in order else len(order)
        if child_idx > new_idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


def qn_local(tag):
    return tag.split("}")[-1] if "}" in tag else tag


def _shade(tcPr, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_ordered(tcPr, shd, _TCPR_ORDER)


def _no_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    _insert_ordered(table._tbl.tblPr, borders, _TBLPR_ORDER)


def _fix_zoom(doc):
    # python-docx's own default settings.xml emits <w:zoom w:val="bestFit"/>
    # without the required w:percent attribute -- present even in a totally
    # empty document with none of our content, so it's a library default, not
    # something our table-building code causes. Still, patch it for a clean
    # validation: Word tolerates it, but no reason to leave an avoidable
    # schema violation in a file we're specifically trying to make bulletproof.
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def _cell_text(cell, text, bold=False):
    # Cell already has one empty paragraph by default; reuse it instead of
    # leaving a stray blank paragraph before the real content.
    p = cell.paragraphs[0]
    lines = str(text).split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()  # soft line break, stays in the same cell/paragraph
        run = p.add_run(line)
        run.bold = bold


def _orange_bar(doc, col_widths):
    table = doc.add_table(rows=1, cols=len(col_widths))
    table.autofit = False
    _no_borders(table)
    for c, w in zip(table.columns, col_widths):
        c.width = Inches(w)
    cell = table.cell(0, 0)
    for i in range(1, len(col_widths)):
        cell = cell.merge(table.cell(0, i))
    _shade(cell._tc.get_or_add_tcPr(), ORANGE_FILL)
    cell.paragraphs[0].add_run(" ").font.size = Pt(4)
    return table


def _header_bar(doc, label, col_widths):
    table = doc.add_table(rows=1, cols=len(col_widths))
    table.autofit = False
    _no_borders(table)
    for c, w in zip(table.columns, col_widths):
        c.width = Inches(w)
    cell = table.cell(0, 0)
    for i in range(1, len(col_widths)):
        cell = cell.merge(table.cell(0, i))
    _shade(cell._tc.get_or_add_tcPr(), LIGHT_BLUE_FILL)
    _cell_text(cell, label, bold=True)
    return table


def _label_value_table(doc, rows, col_widths):
    table = doc.add_table(rows=len(rows), cols=len(col_widths))
    table.autofit = False
    _no_borders(table)
    for c, w in zip(table.columns, col_widths):
        c.width = Inches(w)
    for r, cells in enumerate(rows):
        for c, text in enumerate(cells):
            cell = table.cell(r, c)
            _shade(cell._tc.get_or_add_tcPr(), LIGHT_BLUE_FILL)
            _cell_text(cell, text, bold=(c == 0))
    return table


def generate_docx(row: dict, path: Path):
    body = row.get("device", "")

    # Reuse the CAD/Config detection already done correctly in parse_lead()
    # (row["LeadComments"], derived from the full untruncated body) rather
    # than re-deriving it from row["device"], which extract_device() already
    # truncated to start exactly at "Distributor:".
    intro_line2 = (
        "Please find the following new RFQ and URGENT lead from Dorner Config and process accordingly."
        if row.get("LeadComments") == CONFIG_COMMENT
        else "Please find the following new RFQ and URGENT lead from Dorner CAD and process accordingly."
    )

    doc = Document()
    _fix_zoom(doc)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.5)
    sec.top_margin = sec.bottom_margin = Inches(0.5)

    doc.add_paragraph("Dorner Distributor,")
    doc.add_paragraph(intro_line2)
    doc.add_paragraph("")

    LABEL_W, VALUE_W = 1.2, 2.27  # sums to ~3.47in, matching the original narrow left-side layout

    _orange_bar(doc, [LABEL_W, VALUE_W])
    _label_value_table(doc, [["Distributor:", first_match(body, [r"Distributor:[ \t]*([^\n]*)"]) or row.get("Referral", "")]], [LABEL_W, VALUE_W])
    doc.add_paragraph("")

    _orange_bar(doc, [LABEL_W, VALUE_W])
    _header_bar(doc, "Customer Contact Info:", [LABEL_W, VALUE_W])
    _label_value_table(doc, [
        ["Name:", (row.get("FirstName", "") + " " + row.get("LastName", "")).strip()],
        ["Title:", row.get("ContactTitle", "")],
        ["Industry:", first_match(body, [r"Industry:\s*(.*?)(?:\nCompany:)"])],
        ["Company:", row.get("Company", "")],
        ["Phone:", row.get("PhoneSupplied", "")],
        ["Email:", row.get("Email", "")],
    ], [LABEL_W, VALUE_W])
    doc.add_paragraph("")

    _orange_bar(doc, [LABEL_W, VALUE_W])
    _header_bar(doc, "Customer Contact Info:", [LABEL_W, VALUE_W])
    addr = "\n".join([x for x in [row.get("Address", ""), " ".join([row.get("City", ""), row.get("State", ""), row.get("ZipCode", ""), row.get("Country", "")]).strip()] if x])
    _label_value_table(doc, [["Address:", addr]], [LABEL_W, VALUE_W])
    doc.add_paragraph("")

    QTY_W = [1.5, 2.32, 2.5]  # sums to ~6.32in, matching the original wider quote/total row
    _orange_bar(doc, QTY_W)
    qtable = doc.add_table(rows=1, cols=3)
    qtable.autofit = False
    _no_borders(qtable)
    for c, w in zip(qtable.columns, QTY_W):
        c.width = Inches(w)
    _shade(qtable.cell(0, 0)._tc.get_or_add_tcPr(), LIGHT_BLUE_FILL)
    _shade(qtable.cell(0, 1)._tc.get_or_add_tcPr(), LIGHT_BLUE_FILL)
    _shade(qtable.cell(0, 2)._tc.get_or_add_tcPr(), LIGHT_BLUE_FILL)
    p0 = qtable.cell(0, 0).paragraphs[0]
    p0.add_run("Dorner Quote: ").bold = True
    p0.add_run(str(row.get("_quote", "")))
    qtable.cell(0, 1).paragraphs[0].add_run("Grand Total:").bold = True
    p2 = qtable.cell(0, 2).paragraphs[0]
    r2 = p2.add_run(str(row.get("GrandTotal", "")))
    r2.bold = True
    r2.underline = True
    doc.add_paragraph("")

    # Plain quote/pricing/notes text, unstyled -- same content as before.
    qpos = body.find("Dorner Quote:")
    detail = body[qpos:] if qpos >= 0 else body
    for line in detail.split("\n"):
        doc.add_paragraph(line)

    doc.save(str(path))


def generate_pdf(row: dict, path: Path):
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.45*inch, leftMargin=0.45*inch, topMargin=0.45*inch, bottomMargin=0.45*inch)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("normal", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12)
    bold = ParagraphStyle("bold", parent=normal, fontName="Helvetica-Bold")
    story = []

    def p(text, style=normal):
        story.append(Paragraph(html.escape(str(text)).replace("\n", "<br/>").replace("  ", "&nbsp;&nbsp;"), style))

    intro_end = row.get("device", "").find("Distributor:")
    intro = row.get("device", "")[:intro_end].strip() if intro_end > 0 else "Dorner Distributor,"
    p(intro)
    story.append(Spacer(1, 12))

    def info_table(rows, widths=(1.4*inch, 3.2*inch)):
        data = [[Paragraph(f"<b>{html.escape(a)}</b>", normal), Paragraph(html.escape(str(b)).replace("\n", "<br/>") , normal)] for a,b in rows]
        t = Table(data, colWidths=list(widths))
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#c6d9f1")),
            ("BOX", (0,0), (-1,-1), 0.25, colors.white),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 6), ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 5), ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]))
        bar = Table([[""]], colWidths=[sum(widths)], rowHeights=[18])
        bar.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#ff6600"))]))
        story.extend([bar, t, Spacer(1, 14)])

    info_table([("Distributor:", first_match(row.get("device", ""), [r"Distributor:\s*([^\n]+)"]))])
    info_table([
        ("Customer Contact Info:", ""),
        ("Name:", (row.get("FirstName", "") + " " + row.get("LastName", "")).strip()),
        ("Title:", row.get("ContactTitle", "")),
        ("Industry:", first_match(row.get("device", ""), [r"Industry:\s*(.*?)(?:\nCompany:)"])),
        ("Company:", row.get("Company", "")),
        ("Phone:", row.get("PhoneSupplied", "")),
        ("Email:", row.get("Email", "")),
    ])
    info_table([("Address:", "\n".join([row.get("Address", ""), " ".join([row.get("City", ""), row.get("State", ""), row.get("ZipCode", ""), row.get("Country", "")]).strip()]))])
    info_table([("Dorner Quote:", row.get("_quote", "")), ("Grand Total:", row.get("GrandTotal", ""))], widths=(2*inch, 4.5*inch))

    detail = row.get("device", "")
    qpos = detail.find("Dorner Quote:")
    if qpos >= 0:
        detail = detail[qpos:]
    p(detail)
    doc.build(story)


def build_excel(rows: list[dict]) -> bytes:
    clean_rows = []
    for r in rows:
        clean_rows.append({h: r.get(h, "") for h in HEADERS})
    df = pd.DataFrame(clean_rows, columns=HEADERS)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Leads")
    bio.seek(0)
    wb = load_workbook(bio)
    ws = wb["Leads"]
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for idx, col in enumerate(ws.columns, start=1):
        max_len = 12
        for cell in col:
            val = str(cell.value or "")
            max_len = max(max_len, min(len(val), 40))
        ws.column_dimensions[get_column_letter(idx)].width = max_len + 2
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def process_files(uploaded_files):
    rows = []
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
        for up in uploaded_files:
            data = up.read()
            subject, body, raw_date, sender = parse_msg_bytes(data, up.name)
            row = parse_lead(subject, body, raw_date, up.name)
            rows.append(row)
            base = row["_file_base"]
            tmp_pdf = Path("/tmp") / f"{base}.pdf"
            tmp_doc = Path("/tmp") / f"{base}.docx"
            generate_pdf(row, tmp_pdf)
            generate_docx(row, tmp_doc)
            z.writestr(f"{base}.msg", data)
            z.write(tmp_pdf, f"{base}.pdf")
            z.write(tmp_doc, f"{base}.docx")
        excel_bytes = build_excel(rows)
        z.writestr("Dorner_Leads_Output.xlsx", excel_bytes)
    return rows, zip_buf.getvalue(), build_excel(rows)


st.set_page_config(page_title=APP_TITLE, layout="wide")
st.title(APP_TITLE)
st.caption("Upload one or many Dorner .msg/.txt leads. App creates styled PDF, native Word .docx, MSG copies, and one Excel output with all rows.")
files = st.file_uploader("Upload Dorner lead file(s)", type=["msg", "txt", "eml"], accept_multiple_files=True)

if files:
    try:
        rows, zip_bytes, excel_bytes = process_files(files)
        st.success(f"Parsed {len(rows)} lead(s). Excel will contain {len(rows)} row(s).")
        preview_cols = ["Brand", "Product", "ReceivedDateTime", "FirstName", "LastName", "Company", "Email", "PhoneSupplied", "Keyword", "GrandTotal", "PDF"]
        st.dataframe(pd.DataFrame([{c: r.get(c, "") for c in preview_cols} for r in rows]), use_container_width=True)
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("Download Excel (all rows)", data=excel_bytes, file_name="Dorner_Leads_Output.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        with c2:
            st.download_button("Download All ZIP", data=zip_bytes, file_name="Dorner_Output_All.zip", mime="application/zip")
        with st.expander("Device text preview"):
            for r in rows:
                st.subheader(r.get("Company", "Lead"))
                st.text_area("device", r.get("device", ""), height=220, key=r.get("_file_base"))
    except Exception as e:
        st.error(f"Processing failed: {e}")
        st.exception(e)
